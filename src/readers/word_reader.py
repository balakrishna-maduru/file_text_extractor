from pathlib import Path
from typing import Optional, Dict
from docx import Document as DocxDocument


class WordReader:
    """Docx parser."""

    def load_data(
        self,
        file: Path,
        extra_info: Optional[Dict] = None,
        image_conversion_api: Optional[callable] = None,
    ) -> str:
        """Parse file and replace images with converted text."""
        if not isinstance(file, Path):
            file = Path(file)

        doc = DocxDocument(file)
        full_text = []

        # Handling images in the document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_location = rel.target_part.blob
                image_filename = rel.target_ref.split("/")[-1]

                # Save image locally
                with open(image_filename, "wb") as image_file:
                    image_file.write(image_location)

                if image_conversion_api:
                    extracted_text = image_conversion_api(image_filename)

                    # Find where the image is in the text and replace it
                    for para in doc.paragraphs:
                        if image_filename in para.text:
                            para.text = para.text.replace(image_filename, extracted_text)
                            break  # Assuming one image per paragraph

        # After processing, extract the full text
        for para in doc.paragraphs:
            full_text.append(para.text)

        return "\n".join(full_text)